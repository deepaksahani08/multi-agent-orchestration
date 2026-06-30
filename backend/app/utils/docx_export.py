from io import BytesIO

from docx import Document


def export_docx(markdown: str) -> BytesIO:
    """
    Convert Markdown to DOCX.

    Supported Markdown:

    # Heading 1
    ## Heading 2
    ### Heading 3
    - Bullet List
    Paragraph
    """

    document = Document()

    for line in markdown.splitlines():

        line = line.strip()

        if not line:
            continue

        if line.startswith("# "):
            document.add_heading(line[2:], level=1)

        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)

        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)

        elif line.startswith("- "):
            document.add_paragraph(
                line[2:],
                style="List Bullet",
            )

        else:
            document.add_paragraph(line)

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer